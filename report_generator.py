"""
Report Generator Module for NGO Data Helpers
Handles PDF and DOCX generation from report data
"""

import io
from datetime import datetime
from typing import Dict, Any, Optional
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import streamlit as st


class ReportGenerator:
    """Generate PDF and DOCX reports from report data"""
    
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        """Setup custom styles for reports"""
        # Custom styles for PDF
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1,  # Center alignment
            textColor=colors.darkblue
        ))
        
        self.styles.add(ParagraphStyle(
            name='CustomHeading',
            parent=self.styles['Heading2'],
            fontSize=16,
            spaceAfter=12,
            textColor=colors.darkblue
        ))
        
        self.styles.add(ParagraphStyle(
            name='CustomBody',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            leftIndent=0
        ))
    
    def generate_pdf(self, report_data: Dict[str, Any], country: str, date_range: Dict[str, str]) -> bytes:
        """
        Generate PDF report
        
        Args:
            report_data: Report data dictionary
            country: Country name
            date_range: Date range dictionary
            
        Returns:
            PDF bytes
        """
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
        
        # Build content
        story = []
        
        # Title
        title = f"NGO Data Helpers Report: {country}"
        story.append(Paragraph(title, self.styles['CustomTitle']))
        story.append(Spacer(1, 12))
        
        # Subtitle
        subtitle = f"Date Range: {date_range['start_date']} to {date_range['end_date']}"
        story.append(Paragraph(subtitle, self.styles['CustomBody']))
        story.append(Spacer(1, 12))
        
        # Generated date
        generated_date = f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
        story.append(Paragraph(generated_date, self.styles['CustomBody']))
        story.append(Spacer(1, 20))
        
        # Executive Summary
        story.append(Paragraph("Executive Summary", self.styles['CustomHeading']))
        story.append(Paragraph(report_data.get('summary', 'No summary available.'), self.styles['CustomBody']))
        story.append(Spacer(1, 20))
        
        # Key Events
        story.append(Paragraph("Key Events", self.styles['CustomHeading']))
        key_events = report_data.get('key_events', [])
        if key_events:
            for i, event in enumerate(key_events, 1):
                story.append(Paragraph(f"{i}. {event}", self.styles['CustomBody']))
        else:
            story.append(Paragraph("No key events available.", self.styles['CustomBody']))
        story.append(Spacer(1, 20))
        
        # Trends
        story.append(Paragraph("Trends", self.styles['CustomHeading']))
        trends = report_data.get('trends', [])
        if trends:
            for i, trend in enumerate(trends, 1):
                story.append(Paragraph(f"{i}. {trend}", self.styles['CustomBody']))
        else:
            story.append(Paragraph("No trends available.", self.styles['CustomBody']))
        story.append(Spacer(1, 20))
        
        # Risks
        story.append(Paragraph("Risks", self.styles['CustomHeading']))
        risks = report_data.get('risks', [])
        if risks:
            for i, risk in enumerate(risks, 1):
                story.append(Paragraph(f"{i}. {risk}", self.styles['CustomBody']))
        else:
            story.append(Paragraph("No risks identified.", self.styles['CustomBody']))
        story.append(Spacer(1, 20))
        
        # Data Sources
        story.append(Paragraph("Data Sources", self.styles['CustomHeading']))
        sources = report_data.get('sources', [])
        if sources:
            for source in sources:
                story.append(Paragraph(f"• {source}", self.styles['CustomBody']))
        else:
            story.append(Paragraph("• ACLED (Armed Conflict Location & Event Data Project)", self.styles['CustomBody']))
            story.append(Paragraph("• GDELT (Global Database of Events, Language, and Tone)", self.styles['CustomBody']))
            story.append(Paragraph("• ReliefWeb (Humanitarian Information Service)", self.styles['CustomBody']))
            story.append(Paragraph("• World Bank Open Data", self.styles['CustomBody']))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    def generate_docx(self, report_data: Dict[str, Any], country: str, date_range: Dict[str, str]) -> bytes:
        """
        Generate DOCX report
        
        Args:
            report_data: Report data dictionary
            country: Country name
            date_range: Date range dictionary
            
        Returns:
            DOCX bytes
        """
        doc = Document()
        
        # Title
        title = doc.add_heading(f"NGO Data Helpers Report: {country}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph(f"Date Range: {date_range['start_date']} to {date_range['end_date']}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Generated date
        generated_date = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        generated_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add some space
        doc.add_paragraph()
        
        # Executive Summary
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(report_data.get('summary', 'No summary available.'))
        doc.add_paragraph()
        
        # Key Events
        doc.add_heading("Key Events", level=1)
        key_events = report_data.get('key_events', [])
        if key_events:
            for i, event in enumerate(key_events, 1):
                doc.add_paragraph(f"{i}. {event}", style='List Number')
        else:
            doc.add_paragraph("No key events available.")
        doc.add_paragraph()
        
        # Trends
        doc.add_heading("Trends", level=1)
        trends = report_data.get('trends', [])
        if trends:
            for i, trend in enumerate(trends, 1):
                doc.add_paragraph(f"{i}. {trend}", style='List Number')
        else:
            doc.add_paragraph("No trends available.")
        doc.add_paragraph()
        
        # Risks
        doc.add_heading("Risks", level=1)
        risks = report_data.get('risks', [])
        if risks:
            for i, risk in enumerate(risks, 1):
                doc.add_paragraph(f"{i}. {risk}", style='List Number')
        else:
            doc.add_paragraph("No risks identified.")
        doc.add_paragraph()
        
        # Data Sources
        doc.add_heading("Data Sources", level=1)
        sources = report_data.get('sources', [])
        if sources:
            for source in sources:
                doc.add_paragraph(f"• {source}")
        else:
            doc.add_paragraph("• ACLED (Armed Conflict Location & Event Data Project)")
            doc.add_paragraph("• GDELT (Global Database of Events, Language, and Tone)")
            doc.add_paragraph("• ReliefWeb (Humanitarian Information Service)")
            doc.add_paragraph("• World Bank Open Data")
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()


def download_pdf_report(report_data: Dict[str, Any], country: str, date_range: Dict[str, str]) -> None:
    """
    Download PDF report using Streamlit
    
    Args:
        report_data: Report data dictionary
        country: Country name
        date_range: Date range dictionary
    """
    try:
        generator = ReportGenerator()
        pdf_bytes = generator.generate_pdf(report_data, country, date_range)
        
        filename = f"NGO_Report_{country.replace(' ', '_')}_{date_range['start_date']}_to_{date_range['end_date']}.pdf"
        
        st.download_button(
            label="📄 Download PDF Report",
            data=pdf_bytes,
            file_name=filename,
            mime="application/pdf",
            type="primary"
        )
    except Exception as e:
        st.error(f"Error generating PDF: {str(e)}")


def download_docx_report(report_data: Dict[str, Any], country: str, date_range: Dict[str, str]) -> None:
    """
    Download DOCX report using Streamlit
    
    Args:
        report_data: Report data dictionary
        country: Country name
        date_range: Date range dictionary
    """
    try:
        generator = ReportGenerator()
        docx_bytes = generator.generate_docx(report_data, country, date_range)
        
        filename = f"NGO_Report_{country.replace(' ', '_')}_{date_range['start_date']}_to_{date_range['end_date']}.docx"
        
        st.download_button(
            label="📝 Download DOCX Report",
            data=docx_bytes,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        st.error(f"Error generating DOCX: {str(e)}")
